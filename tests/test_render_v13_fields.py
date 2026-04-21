"""Tests du render v1.3.1 (nouveaux champs conditionnels).

Valide :
- Bench pré-v1.3 (LISI) : render inchangé (3 onglets xlsx, pas de section "Exclusions")
- Bench v1.3 avec exclus_dealbreaker : onglet "Exclusions" xlsx + section dans docx
- Bench v1.3 avec degraded_mode_disclaimer : encart rouge dans docx
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

from lib._common import BenchContext
from lib._docx import write_docx
from lib._xlsx import write_xlsx
from openpyxl import load_workbook
from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
LISI_BENCH = REPO_ROOT / "examples" / "generic-demo" / "bench.json"


# ---------------------------------------------------------------------------
# Non-régression bench pré-v1.3 (LISI)
# ---------------------------------------------------------------------------


@pytest.fixture
def lisi_ctx():
    return BenchContext.load(LISI_BENCH, None)


def test_lisi_xlsx_has_no_exclusions_sheet(lisi_ctx, tmp_path):
    out = tmp_path / "lisi.xlsx"
    write_xlsx(lisi_ctx, out)
    wb = load_workbook(out)
    assert "Exclusions" not in wb.sheetnames, (
        "LISI bench n'a pas exclus_dealbreaker → onglet Exclusions ne doit pas être créé"
    )
    # Onglets standard attendus
    assert set(wb.sheetnames) == {"Matrice", "Détails", "Sources"}


def test_lisi_docx_has_no_degraded_disclaimer(lisi_ctx, tmp_path):
    out = tmp_path / "lisi.docx"
    write_docx(lisi_ctx, out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    # Le texte MODE DEGRADE n'apparaît jamais sur LISI
    assert "MODE DEGRADE" not in full_text


# ---------------------------------------------------------------------------
# Bench v1.3 avec nouveaux champs
# ---------------------------------------------------------------------------


def _make_v13_bench_minimal() -> dict:
    """Construit un bench minimal v1.3 pour tests de render."""
    return {
        "meta": {
            "client": "Test Client",
            "mission": "Test mission v1.3",
            "date": "2026-04-21",
            "version": "1.0",
            "mode": "standard",
            "consultant": "Test User",
            "degraded_threshold": 0.5,
        },
        "grid_ref": "scoring-grids/lisi-souverainete.json",
        "exec_summary": {
            "context": "Bench test v1.3 avec exclusions + disclaimer",
            "three_takeaways": [
                "Takeaway 1 test",
                "Takeaway 2 test",
                "Takeaway 3 test",
            ],
            "exclus_dealbreaker": ["vendor-x"],
            "degraded_mode_disclaimer": (
                "Ce bench a ete produit en MODE DEGRADE. 8/10 fiches en partial et 2/10 "
                "en shallow. Ratio pondere (shallow+partial/2)/total=0.60 > seuil 0.50."
            ),
            "ranking": [
                {"rank": 1, "actor_id": "vendor-a", "weighted_score": 4.5, "positioning": "Top"},
            ],
        },
        "actors": [
            {
                "id": "vendor-a",
                "name": "Vendor A",
                "weighted_score": 4.5,
                "rank": 1,
                "dealbreaker_verdict": "pass",
                "investigation_depth": "full",
                "scoring": [
                    {"criterion_id": "maturite_demarche", "score": 4, "weight": 0.20, "justification": "ok"},
                    {"criterion_id": "independance_stack", "score": 5, "weight": 0.25, "justification": "ok"},
                    {"criterion_id": "ampleur_investissement", "score": 4, "weight": 0.15, "justification": "ok"},
                    {"criterion_id": "pertinence_lisi", "score": 5, "weight": 0.15, "justification": "ok"},
                    {"criterion_id": "approche_by_design_vs_remediation", "score": 4, "weight": 0.10, "justification": "ok"},
                    {"criterion_id": "traction_resultats", "score": 4, "weight": 0.15, "justification": "ok"},
                ],
            },
            {
                "id": "vendor-x",
                "name": "Vendor X (exclu)",
                "weighted_score": 0,
                "rank": 99,
                "dealbreaker_verdict": "violated",
                "dealbreaker_violations": [
                    {
                        "rule_id": "db_hosting_eu",
                        "evidence": "Hebergement AWS us-east-1 exclusivement (page architecture 2026-04-15)",
                        "source_id": "vendor-x-1",
                    }
                ],
                "investigation_depth": "partial",
                "scoring": [
                    {"criterion_id": "maturite_demarche", "score": 3, "weight": 0.20, "justification": "n/a"},
                    {"criterion_id": "independance_stack", "score": 1, "weight": 0.25, "justification": "n/a"},
                    {"criterion_id": "ampleur_investissement", "score": 3, "weight": 0.15, "justification": "n/a"},
                    {"criterion_id": "pertinence_lisi", "score": 3, "weight": 0.15, "justification": "n/a"},
                    {"criterion_id": "approche_by_design_vs_remediation", "score": 2, "weight": 0.10, "justification": "n/a"},
                    {"criterion_id": "traction_resultats", "score": 3, "weight": 0.15, "justification": "n/a"},
                ],
            },
        ],
        "sources": [
            {
                "id": "vendor-x-1",
                "actor_ref": "vendor-x",
                "title": "Page architecture Vendor X",
                "date": "2026-04-15",
                "url": "https://vendor-x.example.com/architecture",
                "type": "primary",
                "reliability": 4,
            },
        ],
    }


GRID_PATH = REPO_ROOT / "scoring-grids" / "lisi-souverainete.json"


@pytest.fixture
def v13_bench_file(tmp_path):
    bench = _make_v13_bench_minimal()
    p = tmp_path / "bench-v13.json"
    p.write_text(json.dumps(bench), encoding="utf-8")
    return p


def test_v13_bench_xlsx_has_exclusions_sheet(v13_bench_file, tmp_path):
    ctx = BenchContext.load(v13_bench_file, GRID_PATH)
    out = tmp_path / "v13.xlsx"
    write_xlsx(ctx, out)
    wb = load_workbook(out)
    assert "Exclusions" in wb.sheetnames
    ws = wb["Exclusions"]
    # Titre, sous-titre, header, au moins 1 ligne data
    assert ws.max_row >= 5
    # Header attendu
    assert ws.cell(row=4, column=1).value == "Acteur"
    assert ws.cell(row=4, column=2).value == "Règle violée"
    # 1ère ligne data = Vendor X
    assert "Vendor X" in str(ws.cell(row=5, column=1).value)
    assert "db_hosting_eu" in str(ws.cell(row=5, column=2).value)


def test_v13_bench_docx_has_degraded_disclaimer_callout(v13_bench_file, tmp_path):
    ctx = BenchContext.load(v13_bench_file, GRID_PATH)
    out = tmp_path / "v13.docx"
    write_docx(ctx, out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "MODE DEGRADE" in full_text
    assert "RE-TRIANGULATION" in full_text


def test_v13_bench_docx_has_exclusions_section(v13_bench_file, tmp_path):
    ctx = BenchContext.load(v13_bench_file, GRID_PATH)
    out = tmp_path / "v13.docx"
    write_docx(ctx, out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "Exclusions" in full_text
    assert "Vendor X" in full_text
    assert "db_hosting_eu" in full_text
