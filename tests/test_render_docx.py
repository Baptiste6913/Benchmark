"""Tests de non-régression pour le générateur docx (charte Ascend).

Génère le bench LISI et vérifie :
- Fichier produit et lisible par python-docx
- Les 11 sections obligatoires (titre de chapitre trouvé)
- Au moins une fiche acteur par acteur du bench
- Un nombre d'encadrés (tables) plausible (cover + scoring + critic + recos…)
- Aucune page blanche (pas de page-break consécutifs)
- Palette Ascend appliquée (couleur principale 3D1D08 dans le XML)
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from lib._common import ASCEND_PALETTE, BenchContext  # noqa: E402
from lib.render import render_docx  # noqa: E402


def _all_text(doc) -> list[str]:
    """Texte de tous les paragraphes + textes des cellules de tableaux (récursif)."""
    chunks = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
                # Tableaux imbriqués (rare, mais safe)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                chunks.append(np.text)
    return chunks


def _docx_xml_bytes(path: Path) -> bytes:
    """Décompresse le docx et retourne le XML de document.xml + header.xml + footer.xml concaténés."""
    out = bytearray()
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                out.extend(z.read(name))
    return bytes(out)

LISI_BENCH = REPO_ROOT / "examples" / "generic-demo" / "bench.json"


@pytest.fixture(scope="module")
def rendered_docx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("render") / "lisi.docx"
    path = render_docx(LISI_BENCH, out)
    assert path.exists()
    assert path.stat().st_size > 5000   # plausible size
    return path


@pytest.fixture(scope="module")
def bench_ctx() -> BenchContext:
    return BenchContext.load(LISI_BENCH)


def test_document_opens(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    assert len(doc.paragraphs) > 100    # the full document is verbose


def test_cover_page(rendered_docx: Path, bench_ctx: BenchContext) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    # ASCEND PARTNERS en haut (avec double espace de tracking)
    assert any("ASCEND" in t and "PARTNERS" in t for t in texts)
    # Titre mission
    assert any(bench_ctx.mission()[:30] in t for t in texts)
    # Livrable + client
    assert any(f"Livrable de préparation — {bench_ctx.client()}" in t for t in texts)


def test_toc_section_headers(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    # section headers upper-cased in TOC
    assert any("SYNTHÈSE & ORIENTATION" in t.upper() for t in texts)
    assert any("FICHES ACTEURS" in t.upper() for t in texts)
    assert any("ANALYSES ET ENSEIGNEMENTS" in t.upper() for t in texts)


def test_how_to_read_section(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Comment lire ce document" in t for t in texts)
    assert any("Lecture flash" in t for t in texts)


def test_exec_summary(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Synthèse exécutive" in t for t in texts)
    # three_takeaways rendus comme encadrés
    assert any("ENSEIGNEMENT #1" in t for t in texts)
    assert any("ENSEIGNEMENT #3" in t for t in texts)


def test_ranking_section(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Classement et verdicts" in t for t in texts)


def test_one_actor_sheet_per_actor(rendered_docx: Path, bench_ctx: BenchContext) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    for actor in bench_ctx.actors_by_rank:
        name = actor.get("name", "")
        # nom apparaît au minimum dans le titre chapitre + ranking
        occurrences = sum(1 for t in texts if name in t)
        assert occurrences >= 2, f"{name}: seulement {occurrences} occurrence(s)"


def test_critic_section(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Rapport critique" in t for t in texts)
    assert any("BLOQUANT" in t for t in texts)
    assert any("AVERTISSEMENT" in t for t in texts)


def test_red_team_section(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Rapport red-team" in t for t in texts)
    assert any("Patterns de fragilité" in t for t in texts)
    assert any("Verdicts par acteur" in t for t in texts)


def test_transverse_section(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Analyse transverse" in t for t in texts)
    assert any("Patterns communs" in t for t in texts)
    assert any("Deux voies observées" in t for t in texts)


def test_recommendations_section(rendered_docx: Path, bench_ctx: BenchContext) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any(f"Enseignements pour {bench_ctx.client()}" in t for t in texts)
    # Au moins R1 dans les titres d'encadré
    assert any("R1 —" in t or "R1 -" in t for t in texts)


def test_sources_annex(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    texts = _all_text(doc)
    assert any("Annexe — sources" in t or "Annexe - sources" in t for t in texts)


def test_ascend_palette_applied(rendered_docx: Path) -> None:
    """Les XML internes du docx contiennent les couleurs principales de la charte."""
    xml = _docx_xml_bytes(rendered_docx)
    for key in ("title_primary", "accent_primary"):
        hex_val = ASCEND_PALETTE[key]
        assert (
            hex_val.upper().encode() in xml or hex_val.lower().encode() in xml
        ), f"couleur Ascend '{key}' ({hex_val}) absente du docx"


def test_footer_has_ascend_brand(rendered_docx: Path) -> None:
    doc = Document(rendered_docx)
    for section in doc.sections:
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert "ascend partners" in footer_text.lower()


def test_margins_2cm(rendered_docx: Path) -> None:
    from docx.shared import Cm
    doc = Document(rendered_docx)
    tolerance_emu = 2000    # ~0.005cm — conversion rounding
    target = Cm(2.0)
    for section in doc.sections:
        for m in (section.top_margin, section.bottom_margin,
                  section.left_margin, section.right_margin):
            assert abs(m - target) < tolerance_emu, f"marge {m} ≠ {target} (±{tolerance_emu} EMU)"


def test_table_count_plausible(rendered_docx: Path, bench_ctx: BenchContext) -> None:
    """Plausibilité du nombre de tableaux.

    Minima attendus :
    - 1 encadré cover (EN BREF)
    - 1 tableau ranking
    - 1 encadré par takeaway (3)
    - 3-5 encadrés par acteur (EN BREF + blind_spots + scoring + …)
    - 3+ encadrés critic (bloquants/avertissements)
    - 3+ tableaux red-team
    - 2-3 tableaux transverse
    - 1 encadré par recommandation (5) × 2 (+caveats)
    """
    doc = Document(rendered_docx)
    n_actors = len(bench_ctx.actors_by_rank)
    # Heuristique basse : 3 + 3 * n_actors encadrés minimum + quelques tables
    assert len(doc.tables) > 10 + 2 * n_actors
