"""Item 1 du hotfix — détecte les pages blanches orphelines.

Pattern causal : une série d'au moins 2 paragraphes vides (pas de texte,
pas de runs) immédiatement suivis d'un paragraphe ne contenant qu'un
page break. Ce pattern pousse Word à rendre une page quasi-vide (footer
seul) avant la nouvelle section.

Le fix (`_page_break()` helper dans lib/_docx.py) réutilise le dernier
paragraphe vide pour y loger le break, au lieu d'ajouter un paragraphe
supplémentaire.

Ce test échoue sur le code pré-hotfix et passe après.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from lib.render import render_docx  # noqa: E402

LISI_BENCH = REPO_ROOT / "examples" / "generic-demo" / "bench.json"


@pytest.fixture(scope="module")
def rendered_docx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("render") / "lisi.docx"
    return render_docx(LISI_BENCH, out)


def _has_page_break(paragraph) -> bool:
    return 'w:type="page"' in paragraph._element.xml


def _is_empty(paragraph) -> bool:
    """Vrai si le paragraphe n'a aucun contenu (pas de texte, pas de runs
    contenant du texte ou un break)."""
    if paragraph.text.strip():
        return False
    # Un paragraphe avec un page break n'est pas considéré comme vide
    # même s'il n'a pas de texte
    if _has_page_break(paragraph):
        return False
    # Un paragraphe avec des runs contenant du texte non plus
    for run in paragraph.runs:
        if run.text.strip():
            return False
    return True


def test_no_excessive_empty_paragraphs_before_page_break(rendered_docx: Path) -> None:
    """Aucune séquence de 3+ paragraphes vides consécutifs ne précède un page break.

    Seuil à 3 : jusqu'à 2 paragraphes vides sont tolérés pour le spacing
    esthétique (cover, sections aérées). 3+ signale un bug d'accumulation de
    spacers superflus qui provoque une page blanche orpheline.
    """
    doc = Document(rendered_docx)
    paragraphs = list(doc.paragraphs)

    problems: list[str] = []
    for i, p in enumerate(paragraphs):
        if not _has_page_break(p):
            continue
        empty_before = 0
        j = i - 1
        while j >= 0 and _is_empty(paragraphs[j]):
            empty_before += 1
            j -= 1
        if empty_before >= 3:
            problems.append(
                f"index {i}: {empty_before} paragraphes vides avant page break "
                f"(précédé par: '{paragraphs[j].text[:40] if j >= 0 else 'BEGIN'}')"
            )

    assert not problems, (
        "Pattern 'paragraphes vides excessifs + page break' détecté (cause probable de pages blanches) :\n  - "
        + "\n  - ".join(problems)
    )


def test_no_page_break_as_first_element_of_document(rendered_docx: Path) -> None:
    """Le premier paragraphe ne doit pas contenir un page break."""
    doc = Document(rendered_docx)
    if doc.paragraphs:
        assert not _has_page_break(doc.paragraphs[0]), "Premier paragraphe ne doit pas être un page break."


def test_page_breaks_count_plausible(rendered_docx: Path) -> None:
    """Le nombre de page breaks doit être dans la fourchette attendue pour ce bench.

    Pour LISI : cover + TOC + how-to-read + exec summary + ranking + 6 fiches acteurs
    + critic + red-team + transverse + recos + sources = ~16 sections majeures,
    donc 10-20 page breaks.
    """
    doc = Document(rendered_docx)
    count = sum(1 for p in doc.paragraphs if _has_page_break(p))
    assert 10 <= count <= 25, f"Nombre de page breaks = {count} hors fourchette attendue 10-25"
