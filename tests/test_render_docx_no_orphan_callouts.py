"""Item 2 du hotfix — détecte les callouts orphelins.

Un callout (encadré Ascend = tableau 1×1 shadé `#FBF1ED` ou `#EEEBE3`) orphelin
apparaît seul en haut d'une page, détaché de son contexte. Causes :
- les rows du callout peuvent se couper entre pages (pas de `cantSplit`)
- le paragraphe qui précède le callout n'a pas `keep_with_next` — Word peut
  séparer le contexte et le callout

Fix (lib/_docx.py) : `_set_row_cant_split` + `_set_paragraph_keep_next`
appliqués systématiquement dans les helpers `_add_box*`.

Ce test échoue sur le code pré-hotfix et passe après.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from lib._common import ASCEND_PALETTE  # noqa: E402
from lib.render import render_docx  # noqa: E402

LISI_BENCH = REPO_ROOT / "examples" / "generic-demo" / "bench.json"

CALLOUT_FILLS = {ASCEND_PALETTE["box_bg_warm"].upper(), ASCEND_PALETTE["box_bg_neutral"].upper()}


@pytest.fixture(scope="module")
def rendered_docx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("render") / "lisi.docx"
    return render_docx(LISI_BENCH, out)


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return (shd.get(qn("w:fill")) or "").upper()


def _is_callout(tbl) -> bool:
    """Un callout est un tableau dont la 1re cellule a un shading Ascend beige."""
    try:
        first_cell = tbl.rows[0].cells[0]
    except IndexError:
        return False
    return _cell_fill(first_cell) in CALLOUT_FILLS


def _row_has_cant_split(row) -> bool:
    tr_pr = row._tr.find(qn("w:trPr"))
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:cantSplit")) is not None


def test_all_callouts_have_cant_split(rendered_docx: Path) -> None:
    """Chaque row d'un callout doit avoir <w:cantSplit/> pour éviter la coupure."""
    doc = Document(rendered_docx)
    callouts = [tbl for tbl in doc.tables if _is_callout(tbl)]
    assert callouts, "Aucun callout détecté — le test est peut-être cassé."

    problems = []
    for i, tbl in enumerate(callouts):
        for j, row in enumerate(tbl.rows):
            if not _row_has_cant_split(row):
                fill = _cell_fill(row.cells[0])
                problems.append(f"callout #{i} (fill={fill}) row {j} sans cantSplit")

    assert not problems, "Callouts sans cantSplit (risque de coupure inter-page) :\n  - " + "\n  - ".join(problems)


def _paragraph_has_keep_next(paragraph) -> bool:
    """Lit `<w:keepNext/>` dans <w:pPr>."""
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return p_pr.find(qn("w:keepNext")) is not None


def test_paragraph_before_callout_has_keep_next(rendered_docx: Path) -> None:
    """Le paragraphe juste avant chaque callout doit être en keep_with_next.

    Cela évite que Word orphanise le callout en haut d'une nouvelle page en le
    séparant de son contexte. Exception : si le callout est le tout premier
    élément du body (pas de paragraphe précédent).
    """
    doc = Document(rendered_docx)
    body = doc.element.body
    problems = []

    # Itérer sur les éléments dans l'ordre du body
    children = list(body.iterchildren())
    for i, child in enumerate(children):
        if not child.tag.endswith("}tbl"):
            continue

        # Vérifier si c'est un callout
        from docx.table import Table
        tbl = Table(child, doc)
        if not _is_callout(tbl):
            continue

        # Chercher le paragraphe précédent non-vide
        j = i - 1
        prev_para = None
        while j >= 0:
            prev = children[j]
            if prev.tag.endswith("}p"):
                from docx.text.paragraph import Paragraph
                para = Paragraph(prev, doc)
                if para.text.strip() or any(r.text.strip() for r in para.runs):
                    prev_para = para
                    break
                # Paragraphe vide → on cherche plus haut (spacer)
            j -= 1

        if prev_para is None:
            continue   # callout en tête de doc, pas de paragraphe avant à attirer

        if not _paragraph_has_keep_next(prev_para):
            problems.append(
                f"Callout après '{prev_para.text[:60]}' — paragraphe précédent sans keep_with_next"
            )

    assert not problems, (
        "Callouts potentiellement orphelins (paragraphe précédent sans keep_with_next) :\n  - "
        + "\n  - ".join(problems)
    )
