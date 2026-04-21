"""Item 3 du hotfix — vérifie que les libellés de critères des livrables
sont les formes accentuées françaises (`label_fr`), pas les IDs ASCII.

Ce test échoue sur le code pré-hotfix (le générateur lit `name` qui est
non-accentué) et passe après (le générateur lit `label_fr` avec fallback
sur `name`).
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from lib.render import render_docx, render_xlsx  # noqa: E402

LISI_BENCH = REPO_ROOT / "examples" / "generic-demo" / "bench.json"

# Noms de critères NON-accentués de la grille LISI (champs `name`).
# Doivent DISPARAÎTRE des HEADERS des livrables (matrice xlsx row 5, tableaux
# scoring docx). Peuvent rester dans des justifications narratives — hors scope
# de ce test.
NON_ACCENTED_HEADERS = [
    "Maturite de la demarche souverainete",
    "Independance stack technique (infra / modeles / donnees)",
    "Pertinence pour un industriel sous-traitant defense/aero",
    "Traction / resultats demontres",
]

# Formes accentuées attendues (valeurs de `label_fr`)
ACCENTED_HEADERS = [
    "Maturité de la démarche de souveraineté",
    "Indépendance du stack (infra / modèles / données)",
    "Pertinence pour un industriel sous-traitant défense/aéro",
    "Traction et résultats démontrés",
]


@pytest.fixture(scope="module")
def rendered_xlsx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("render") / "lisi.xlsx"
    return render_xlsx(LISI_BENCH, out)


@pytest.fixture(scope="module")
def rendered_docx(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("render") / "lisi.docx"
    return render_docx(LISI_BENCH, out)


def _all_cells_text(rendered_xlsx: Path) -> str:
    wb = load_workbook(rendered_xlsx)
    chunks = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            for val in row:
                if isinstance(val, str):
                    chunks.append(val)
    return "\n".join(chunks)


def _all_docx_text(rendered_docx: Path) -> str:
    """Concatène le texte de tous les paragraphes + cellules de tables du docx."""
    doc = Document(rendered_docx)
    chunks = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


def _docx_raw_xml(rendered_docx: Path) -> str:
    """Contenu XML brut du docx (décompressé) pour la recherche de formes exactes."""
    parts = []
    with zipfile.ZipFile(rendered_docx) as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                parts.append(z.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(parts)


def test_xlsx_matrice_headers_accented(rendered_xlsx: Path) -> None:
    """La ligne 5 (headers de colonnes) de l'onglet Matrice doit utiliser label_fr."""
    wb = load_workbook(rendered_xlsx)
    ws = wb["Matrice"]
    headers = [ws.cell(row=5, column=c).value for c in range(1, ws.max_column + 1)]
    headers_joined = " | ".join(str(h) for h in headers if h)

    for form in NON_ACCENTED_HEADERS:
        assert form not in headers_joined, (
            f"Matrice header contient '{form}' (non-accentué). "
            f"Attendu : le générateur lit `label_fr`. Headers actuels : {headers_joined}"
        )
    for form in ACCENTED_HEADERS:
        assert form in headers_joined, (
            f"Matrice header devrait contenir '{form}'. Headers actuels : {headers_joined}"
        )


def test_xlsx_details_sheet_accented(rendered_xlsx: Path) -> None:
    """L'onglet Détails liste les critères ; chaque ligne doit porter le label_fr."""
    wb = load_workbook(rendered_xlsx)
    ws = wb["Détails"]
    criterion_cells = [ws.cell(row=r, column=2).value for r in range(5, ws.max_row + 1)]
    criterion_joined = " | ".join(str(c) for c in criterion_cells if c)

    for form in NON_ACCENTED_HEADERS:
        assert form not in criterion_joined, f"Détails contient '{form}' (non-accentué)"
    for form in ACCENTED_HEADERS:
        assert form in criterion_joined, f"Détails devrait contenir '{form}'"


def test_docx_uses_accented_labels_somewhere(rendered_docx: Path) -> None:
    """Les libellés accentués doivent apparaître dans le docx (au moins dans
    les tableaux de scoring des fiches acteurs)."""
    text = _all_docx_text(rendered_docx)
    for form in ACCENTED_HEADERS:
        assert form in text, f"Le docx devrait contenir la forme accentuée '{form}'."


def test_docx_no_non_accented_headers_in_xml(rendered_docx: Path) -> None:
    """Les noms NON-accentués ne doivent pas apparaître dans le XML du docx.

    Toutes les occurrences de critères en affichage passent par `label_fr`.
    Les `name` non-accentués restent dans le bench.json (justifications)
    mais comme ce test vérifie les libellés de critères exacts, ils ne
    devraient pas y apparaître."""
    xml = _docx_raw_xml(rendered_docx)
    for form in NON_ACCENTED_HEADERS:
        assert form not in xml, (
            f"Forme non-accentuée '{form}' présente dans le docx — "
            "le générateur lit encore `name` au lieu de `label_fr`."
        )
