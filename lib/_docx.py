"""Générateur docx conforme à la charte Ascend (stricte).

Produit le livrable Word à partir d'un BenchContext (bench.json + grille).
Respecte :
- palette couleurs hex (`ASCEND_PALETTE`)
- polices Signifier/Söhne avec fallback Georgia/Calibri (voir USE_PRIMARY_FONTS)
- structure 11 sections obligatoires (cover, TOC, how-to-read, exec summary,
  ranking, actor sheets, critic, red-team, transverse, recommendations, sources)
- marges 2cm partout, pied de page standard
- encadrés via tableaux 1×1 avec fond coloré (python-docx n'offre pas mieux)
- tableaux scoring/chiffres-clés en-tête `#3D1D08` texte blanc, première colonne
  `#EEEBE3`, bordures `#CBB57E`
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from lib._common import ASCEND_FONTS, ASCEND_PALETTE, BenchContext, criterion_display, verdict_palette

# -----------------------------------------------------------------------------
# Configuration polices
# -----------------------------------------------------------------------------

# Par défaut, on utilise les fallbacks système (Georgia / Calibri) pour que le
# docx s'affiche correctement sur tout poste Windows / macOS. Si Signifier et
# Söhne sont installées (postes consultants Ascend Partners avec la charte
# polices), basculer USE_PRIMARY_FONTS à True ou définir ASCEND_USE_PRIMARY_FONTS
# comme variable d'environnement à "1".
import os

USE_PRIMARY_FONTS = os.environ.get("ASCEND_USE_PRIMARY_FONTS", "0") == "1"

if USE_PRIMARY_FONTS:
    FONT_TITLE = ASCEND_FONTS["title"][0]    # Signifier
    FONT_BODY = ASCEND_FONTS["body"][0]      # Söhne
else:
    FONT_TITLE = ASCEND_FONTS["title"][1]    # Georgia
    FONT_BODY = ASCEND_FONTS["body"][1]      # Calibri


# -----------------------------------------------------------------------------
# Helpers OOXML
# -----------------------------------------------------------------------------


def _set_run(
    run,
    *,
    font: str,
    size_pt: float,
    color_hex: str,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Configure un run avec police + taille + couleur + style."""
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.bold = bold
    run.font.italic = italic
    # Forcer la police sur ascii et hAnsi pour compat cross-platform
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str, size: str = "4") -> None:
    """Met des bordures sur les 4 côtés d'une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:color"), color_hex)
        borders.append(e)
    tc_pr.append(borders)


def _set_row_cant_split(row) -> None:
    """<w:cantSplit/> sur la row — empêche la coupure du tableau entre pages.

    Appliqué systématiquement sur les callouts pour qu'un encadré ne se
    retrouve jamais à cheval sur deux pages.
    """
    tr = row._tr
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        tr.insert(0, tr_pr)
    # Skip si déjà présent
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_paragraph_keep_next(paragraph) -> None:
    """<w:keepNext/> sur le paragraphe — attire l'élément suivant sur la même page.

    Appliqué sur le paragraphe qui précède un callout pour éviter l'orphelinage
    (callout seul en haut de page sans son contexte).
    """
    paragraph.paragraph_format.keep_with_next = True


def _last_content_paragraph(doc):
    """Retourne le dernier paragraphe direct du doc avec du contenu, ou None."""
    for p in reversed(doc.paragraphs):
        if p.text.strip() or any(r.text.strip() for r in p.runs):
            return p
    return None


# -----------------------------------------------------------------------------
# Primitives d'écriture
# -----------------------------------------------------------------------------


def _add_spacer(doc, n: int = 1) -> None:
    for _ in range(n):
        doc.add_paragraph()


def _page_break(doc) -> None:
    """Ajoute un page break en réutilisant le dernier paragraphe s'il est vide.

    Évite le pattern "paragraphes vides consécutifs + paragraphe-break" qui
    génère une page blanche orpheline (bug identifié en hotfix Sprint 2, item 1).
    Si le dernier paragraphe a déjà du contenu ou n'existe pas, retombe sur
    le comportement standard `doc.add_page_break()`.
    """
    if not doc.paragraphs:
        doc.add_page_break()
        return
    last = doc.paragraphs[-1]
    if last.text.strip() or any(r.text.strip() for r in last.runs):
        doc.add_page_break()
        return
    # Paragraphe vide — on y met le break
    if last.runs:
        last.runs[-1].add_break(WD_BREAK.PAGE)
    else:
        run = last.add_run()
        run.add_break(WD_BREAK.PAGE)


def _add_styled_para(doc_or_cell, text: str, *, font: str, size: float, color: str,
                     bold: bool = False, italic: bool = False,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    r = p.add_run(text)
    _set_run(r, font=font, size_pt=size, color_hex=color, bold=bold, italic=italic)


def _add_chapter_title(doc, number: str, title: str) -> None:
    """Titre de chapitre Ascend : numéro rouge brique + titre brun foncé, Signifier/Georgia 20pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"{number}. ")
    _set_run(r1, font=FONT_TITLE, size_pt=20, color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
    r2 = p.add_run(title)
    _set_run(r2, font=FONT_TITLE, size_pt=20, color_hex=ASCEND_PALETTE["title_primary"], bold=True)


def _add_section_header(doc, text: str) -> None:
    """Section header type 'FICHES ACTEURS' — rouge brique, tracking large, 10pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    _set_run(r, font=FONT_BODY, size_pt=10, color_hex=ASCEND_PALETTE["accent_primary"], bold=True)


def _add_subtitle(doc_or_cell, text: str) -> None:
    """Sous-titre type 'Contexte et déclencheur' — 14pt brun foncé."""
    p = doc_or_cell.add_paragraph()
    r = p.add_run(text)
    _set_run(r, font=FONT_TITLE, size_pt=14, color_hex=ASCEND_PALETTE["title_primary"], bold=True)


def _add_box(doc, title: str, lines: list[str], *, bg_color: str,
             title_color: str = None) -> None:
    """Ajoute un encadré Ascend (tableau 1×1) avec titre + corps.

    `lines` : liste de paragraphes. Pour formatter un (label, value), préférer
    `_add_box_kv`.

    Garde-fous anti-orphelin (hotfix item 2) :
    - `cantSplit` sur chaque row → le callout ne se coupe pas entre pages.
    - `keep_with_next` sur le paragraphe précédent → Word garde le callout
      avec son contexte plutôt que de l'orphaniser en haut de page.
    """
    # Attire le callout avec le paragraphe précédent
    prev = _last_content_paragraph(doc)
    if prev is not None:
        _set_paragraph_keep_next(prev)

    title_color = title_color or ASCEND_PALETTE["accent_primary"]
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, bg_color)
    _set_cell_borders(cell, ASCEND_PALETTE["accent_secondary"], size="4")

    # Titre
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    _set_run(r, font=FONT_BODY, size_pt=10, color_hex=title_color, bold=True)

    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        _set_run(r, font=FONT_BODY, size_pt=10, color_hex=ASCEND_PALETTE["title_primary"])

    # Anti-coupure : chaque row reste d'un seul tenant
    for row in tbl.rows:
        _set_row_cant_split(row)


def _add_box_kv(doc, title: str, pairs: list[tuple[str, str]], *, bg_color: str,
                title_color: str = None) -> None:
    """Encadré avec des (label, value) en lignes. Anti-orphelin actif."""
    prev = _last_content_paragraph(doc)
    if prev is not None:
        _set_paragraph_keep_next(prev)

    title_color = title_color or ASCEND_PALETTE["accent_primary"]
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, bg_color)
    _set_cell_borders(cell, ASCEND_PALETTE["accent_secondary"], size="4")

    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    _set_run(r, font=FONT_BODY, size_pt=10, color_hex=title_color, bold=True)

    for label, value in pairs:
        p = cell.add_paragraph()
        r_label = p.add_run(f"{label} — ")
        _set_run(r_label, font=FONT_BODY, size_pt=10,
                 color_hex=ASCEND_PALETTE["title_primary"], bold=True)
        r_val = p.add_run(str(value))
        _set_run(r_val, font=FONT_BODY, size_pt=10, color_hex=ASCEND_PALETTE["title_primary"])

    for row in tbl.rows:
        _set_row_cant_split(row)


def _add_quote_box(doc, quote: str, author: str, date: str) -> None:
    """Encadré citation — fond FBF1ED, titre CITATION en rouge, italique. Anti-orphelin actif."""
    prev = _last_content_paragraph(doc)
    if prev is not None:
        _set_paragraph_keep_next(prev)

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, ASCEND_PALETTE["box_bg_warm"])
    _set_cell_borders(cell, ASCEND_PALETTE["accent_secondary"], size="4")

    p_title = cell.paragraphs[0]
    r = p_title.add_run("CITATION")
    _set_run(r, font=FONT_BODY, size_pt=9, color_hex=ASCEND_PALETTE["accent_primary"], bold=True)

    p_q = cell.add_paragraph()
    r_q = p_q.add_run(f"« {quote} »")
    _set_run(r_q, font=FONT_TITLE, size_pt=12, color_hex=ASCEND_PALETTE["title_primary"], italic=True)

    p_a = cell.add_paragraph()
    r_a = p_a.add_run(f"— {author}, {date}")
    _set_run(r_a, font=FONT_BODY, size_pt=9, color_hex=ASCEND_PALETTE["brun_clair"])

    for row in tbl.rows:
        _set_row_cant_split(row)


def _add_table_header_row(row, values: list[str]) -> None:
    """Style la première ligne d'un tableau (en-tête 3D1D08 blanc)."""
    for cell, value in zip(row.cells, values):
        _set_cell_shading(cell, ASCEND_PALETTE["title_primary"])
        _set_cell_borders(cell, ASCEND_PALETTE["accent_secondary"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        _set_run(r, font=FONT_BODY, size_pt=10, color_hex=ASCEND_PALETTE["white"], bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table_data_row(row, values: list[str], *, highlight_first: bool = True) -> None:
    """Style une ligne data. Premier col = fond EEEBE3, reste = normal."""
    for idx, (cell, value) in enumerate(zip(row.cells, values)):
        _set_cell_borders(cell, ASCEND_PALETTE["accent_secondary"])
        if highlight_first and idx == 0:
            _set_cell_shading(cell, ASCEND_PALETTE["box_bg_neutral"])
        p = cell.paragraphs[0]
        r = p.add_run(str(value) if value is not None else "")
        _set_run(r, font=FONT_BODY, size_pt=9, color_hex=ASCEND_PALETTE["title_primary"],
                 bold=(highlight_first and idx == 0))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


# -----------------------------------------------------------------------------
# Marges + pied de page
# -----------------------------------------------------------------------------


def _setup_sections(doc, ctx: BenchContext) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # Pied de page : "ascend partners" à droite en brun clair, 9pt
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("ascend partners")
        _set_run(r, font=FONT_BODY, size_pt=9, color_hex=ASCEND_PALETTE["brun_clair"])


# -----------------------------------------------------------------------------
# Sections
# -----------------------------------------------------------------------------


def _build_cover(doc, ctx: BenchContext) -> None:
    # "ASCEND PARTNERS" en haut, rouge brique, tracking
    _add_styled_para(
        doc, "ASCEND  PARTNERS",
        font=FONT_BODY, size=10,
        color=ASCEND_PALETTE["accent_primary"], bold=True,
    )

    _add_spacer(doc, 3)

    # Titre principal — titre mission
    _add_styled_para(
        doc, ctx.mission(),
        font=FONT_TITLE, size=28,
        color=ASCEND_PALETTE["title_primary"], bold=True,
    )

    # Sous-titre
    _add_styled_para(
        doc, f"Livrable de préparation — {ctx.client()}",
        font=FONT_BODY, size=13,
        color=ASCEND_PALETTE["brun_clair"], italic=True,
    )

    _add_spacer(doc, 3)

    # Encadré "EN BREF"
    _add_box_kv(
        doc, "EN BREF",
        [
            ("Objet", ctx.mission()),
            ("Panel", f"{len(ctx.actors_by_rank)} acteurs benchés"),
            ("Méthodologie", ctx.bench.get("meta", {}).get("methodology", "")),
            ("Grille de scoring", ctx.grid_path.name),
            ("Mode", ctx.mode()),
            ("Date", ctx.date()),
            ("Consultant", ctx.consultant()),
        ],
        bg_color=ASCEND_PALETTE["box_bg_warm"],
    )

    # Saut de page
    _page_break(doc)


def _build_toc(doc, ctx: BenchContext) -> None:
    _add_styled_para(doc, "Sommaire", font=FONT_TITLE, size=28,
                     color=ASCEND_PALETTE["title_primary"], bold=True)
    _add_spacer(doc)

    _add_section_header(doc, "Synthèse & orientation")
    for num, title in [
        ("01", "Comment lire ce document"),
        ("02", "Synthèse exécutive"),
        ("03", "Classement et verdicts"),
    ]:
        _add_toc_entry(doc, num, title)

    _add_spacer(doc)
    _add_section_header(doc, "Fiches acteurs")
    offset = 4
    for i, actor in enumerate(ctx.actors_by_rank):
        num = f"{offset + i:02d}"
        title = f"{actor.get('name', '')} — {actor.get('weighted_score', 0):.2f} / 5"
        _add_toc_entry(doc, num, title)

    _add_spacer(doc)
    _add_section_header(doc, "Analyses et enseignements")
    offset_next = offset + len(ctx.actors_by_rank)
    for i, title in enumerate([
        "Rapport critique",
        "Rapport red-team",
        "Analyse transverse",
        f"Enseignements pour {ctx.client()}",
        "Annexe — sources",
    ]):
        num = f"{offset_next + i:02d}"
        _add_toc_entry(doc, num, title)

    _page_break(doc)


def _add_toc_entry(doc, number: str, title: str) -> None:
    """Ligne de sommaire : numéro rouge, titre brun foncé, points de fuite."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r_num = p.add_run(f"{number}  ")
    _set_run(r_num, font=FONT_BODY, size_pt=11, color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
    r_title = p.add_run(title)
    _set_run(r_title, font=FONT_BODY, size_pt=11, color_hex=ASCEND_PALETTE["title_primary"], bold=True)


def _build_how_to_read(doc, ctx: BenchContext) -> None:
    _add_chapter_title(doc, "01", "Comment lire ce document")
    _add_spacer(doc)

    _add_styled_para(
        doc,
        "Ce document est construit pour trois scénarios de lecture.",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["title_primary"],
    )

    _add_styled_para(
        doc, "Lecture flash (5 minutes).",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["accent_primary"], bold=True,
    )
    _add_styled_para(
        doc,
        "Page de garde, synthèse exécutive (02), classement et verdicts (03). Vous savez quoi "
        "retenir pour un arbitrage rapide.",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["title_primary"],
    )
    _add_spacer(doc)

    _add_styled_para(
        doc, "Lecture préparation COPIL (30 minutes).",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["accent_primary"], bold=True,
    )
    _add_styled_para(
        doc,
        "Idem + les fiches acteurs du top 3 + analyse transverse + enseignements. "
        "Les angles morts et verdicts red-team vous protègent de la question piège.",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["title_primary"],
    )
    _add_spacer(doc)

    _add_styled_para(
        doc, "Lecture approfondie (1h30).",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["accent_primary"], bold=True,
    )
    _add_styled_para(
        doc,
        "Tout le document dans l'ordre, annexe sources comprise. À faire avant d'engager "
        "une recommandation contractuelle.",
        font=FONT_BODY, size=11, color=ASCEND_PALETTE["title_primary"],
    )
    _add_spacer(doc)

    _add_subtitle(doc, "Structure d'une fiche acteur")
    _add_styled_para(
        doc,
        "Contexte et déclencheur • Stratégie déployée (infra / modèles / données / org) • "
        "Chiffres-clés • Maturité & approche • Scoring par critère • Angles morts • "
        "Verdict red-team • Trous d'information.",
        font=FONT_BODY, size=10, color=ASCEND_PALETTE["title_primary"],
    )
    _add_spacer(doc)

    _add_subtitle(doc, "Codes couleur")
    _add_styled_para(
        doc,
        "Rouge brique Ascend = verdict REFAIRE ou accent critique. "
        "Beige doré = À NUANCER. Bleu pétrole = TIENT LA ROUTE. "
        "Brun clair = métadonnées et sources secondaires.",
        font=FONT_BODY, size=10, color=ASCEND_PALETTE["title_primary"],
    )

    _page_break(doc)


def _build_degraded_disclaimer_callout(doc, ctx: BenchContext) -> None:
    """Encart rouge au début du livrable si bench en mode dégradé (v1.3.1).

    Conditionnel sur exec_summary.degraded_mode_disclaimer non vide. Posé
    AVANT les enseignements pour que le lecteur voit le flag en premier.
    """
    disclaimer = (ctx.bench.get("exec_summary") or {}).get("degraded_mode_disclaimer")
    if not disclaimer:
        return
    _add_box(
        doc, "/!\\ MODE DEGRADE — RE-TRIANGULATION REQUISE",
        [disclaimer],
        bg_color=ASCEND_PALETTE["box_bg_warm"],
        title_color=ASCEND_PALETTE["accent_primary"],
    )


def _build_exec_summary(doc, ctx: BenchContext) -> None:
    exec_ = ctx.bench.get("exec_summary", {}) or {}

    _add_chapter_title(doc, "02", "Synthèse exécutive")
    _add_spacer(doc)

    # v1.3.1 : disclaimer dégradé en tête si présent
    _build_degraded_disclaimer_callout(doc, ctx)

    context_txt = exec_.get("context", "")
    if context_txt:
        _add_styled_para(doc, context_txt, font=FONT_BODY, size=11,
                         color=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    _add_subtitle(doc, "Trois enseignements structurants")
    _add_spacer(doc)

    takeaways = exec_.get("three_takeaways", []) or []
    for i, t in enumerate(takeaways, start=1):
        _add_box(
            doc, f"ENSEIGNEMENT #{i}", [t],
            bg_color=ASCEND_PALETTE["box_bg_neutral"],
            title_color=ASCEND_PALETTE["title_primary"],
        )

    weak = exec_.get("weak_signals", []) or []
    if weak:
        _add_subtitle(doc, "Signaux faibles à remonter")
        for signal in weak:
            _add_box(
                doc, "SIGNAL", [signal],
                bg_color=ASCEND_PALETTE["box_bg_warm"],
                title_color=ASCEND_PALETTE["accent_primary"],
            )

    _page_break(doc)


def _build_ranking(doc, ctx: BenchContext) -> None:
    _add_chapter_title(doc, "03", "Classement et verdicts")
    _add_spacer(doc)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = True
    _add_table_header_row(tbl.rows[0], ["#", "Acteur", "Score /5", "Verdict red-team"])

    for actor in ctx.actors_by_rank:
        row = tbl.add_row()
        verdict = actor.get("red_team_verdict") or "—"
        pal = verdict_palette(verdict)
        _add_table_data_row(
            row,
            [
                actor.get("rank", ""),
                actor.get("name", ""),
                f"{actor.get('weighted_score', 0):.2f}",
                verdict,
            ],
        )
        # Override couleur verdict sur la dernière cellule
        v_cell = row.cells[3]
        _set_cell_shading(v_cell, pal["bg"])
        v_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(pal["fg"])
        v_cell.paragraphs[0].runs[0].font.bold = True

    _add_spacer(doc)

    # Positionnement en liste dense
    _add_subtitle(doc, "Positionnement synthétique")
    _add_spacer(doc)
    for actor in ctx.actors_by_rank:
        positioning = (
            actor.get("positioning_one_liner")
            or (actor.get("tldr") or {}).get("positioning", "")
        )
        p = doc.add_paragraph()
        r_num = p.add_run(f"#{actor.get('rank', '')}  ")
        _set_run(r_num, font=FONT_BODY, size_pt=10,
                 color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
        r_name = p.add_run(f"{actor.get('name', '')} — ")
        _set_run(r_name, font=FONT_BODY, size_pt=10,
                 color_hex=ASCEND_PALETTE["title_primary"], bold=True)
        r_pos = p.add_run(positioning)
        _set_run(r_pos, font=FONT_BODY, size_pt=10,
                 color_hex=ASCEND_PALETTE["title_primary"])

    _page_break(doc)


def _build_actor_sheet(doc, ctx: BenchContext, actor: dict, chapter_num: int) -> None:
    title = f"{actor.get('name', '')} — {actor.get('weighted_score', 0):.2f} / 5"
    _add_chapter_title(doc, f"{chapter_num:02d}", title)
    _add_spacer(doc)

    tldr = actor.get("tldr") or {}
    _add_box_kv(
        doc, "EN BREF",
        [
            ("Rang", f"#{actor.get('rank', '')}"),
            ("Score pondéré", f"{actor.get('weighted_score', 0):.2f} / 5"),
            ("Verdict red-team", actor.get("red_team_verdict") or "—"),
            ("Positionnement", tldr.get("positioning", actor.get("positioning_one_liner", ""))),
            ("Stack emblématique", tldr.get("stack_emblematique", "")),
        ],
        bg_color=ASCEND_PALETTE["box_bg_warm"],
    )

    # Contexte
    context = actor.get("context", "")
    if context:
        _add_subtitle(doc, "Contexte et déclencheur")
        _add_styled_para(doc, context, font=FONT_BODY, size=10,
                         color=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    # Citation-phare
    fq = actor.get("flagship_quote") or {}
    if fq.get("text"):
        _add_quote_box(doc, fq["text"], fq.get("author", ""), fq.get("date", ""))

    # Stratégie
    strategy = actor.get("strategy") or {}
    if any(strategy.values()):
        _add_subtitle(doc, "Stratégie déployée")
        for label, key in [
            ("Infra / cloud", "infra"),
            ("Modèles IA", "models"),
            ("Données et gouvernance", "data_gov"),
            ("Gouvernance / organisation", "org"),
        ]:
            if strategy.get(key):
                p = doc.add_paragraph()
                r_l = p.add_run(f"{label} — ")
                _set_run(r_l, font=FONT_BODY, size_pt=10,
                         color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
                r_v = p.add_run(strategy[key])
                _set_run(r_v, font=FONT_BODY, size_pt=10,
                         color_hex=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    # Chiffres-clés
    figures = actor.get("key_figures", []) or []
    if figures:
        _add_subtitle(doc, "Chiffres-clés")
        tbl = doc.add_table(rows=1, cols=3)
        _add_table_header_row(tbl.rows[0], ["Métrique", "Valeur", "Date"])
        for fig in figures:
            row = tbl.add_row()
            _add_table_data_row(
                row,
                [fig.get("metric", ""), fig.get("value", ""), fig.get("date", "")],
            )
        _add_spacer(doc)

    # Maturity & approach
    ma = actor.get("maturity_and_approach")
    if ma:
        _add_subtitle(doc, "Maturité et approche")
        _add_styled_para(doc, ma, font=FONT_BODY, size=10,
                         color=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    # Scoring
    scoring = actor.get("scoring", []) or []
    if scoring:
        _add_subtitle(doc, "Scoring par critère")
        tbl = doc.add_table(rows=1, cols=4)
        _add_table_header_row(tbl.rows[0], ["Critère", "Poids", "Score /5", "Justification"])
        for row_data in scoring:
            cid = row_data.get("criterion_id", "")
            criterion = next((c for c in ctx.criteria if c["id"] == cid), None)
            criterion_name = criterion_display(criterion) if criterion else cid
            row = tbl.add_row()
            _add_table_data_row(
                row,
                [
                    criterion_name,
                    f"{int(round(row_data.get('weight', 0) * 100))}%",
                    row_data.get("score", ""),
                    row_data.get("justification", ""),
                ],
            )
        _add_spacer(doc)

    # Angles morts
    blinds = actor.get("blind_spots", []) or []
    if blinds:
        verdict = actor.get("red_team_verdict") or "N/A"
        pal = verdict_palette(verdict)
        bg = ASCEND_PALETTE["box_bg_warm"] if "REFAIRE" in verdict.upper() else ASCEND_PALETTE["box_bg_neutral"]
        _add_subtitle(doc, "Angles morts")
        lines = [f"{b.get('num', i+1)}. {b.get('text', '')}" for i, b in enumerate(blinds)]
        _add_box(doc, "ANGLES MORTS", lines, bg_color=bg, title_color=pal["fg"])

    # Trous d'information
    gaps = actor.get("information_gaps", []) or []
    if gaps:
        _add_subtitle(doc, "Trous d'information")
        for g in gaps:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(f"• {g}")
            _set_run(r, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["brun_clair"])

    _page_break(doc)


def _build_critic(doc, ctx: BenchContext, chapter_num: int) -> None:
    critic = ctx.bench.get("critic") or {}
    if not critic:
        return

    _add_chapter_title(doc, f"{chapter_num:02d}", "Rapport critique")
    _add_spacer(doc)

    blocking = critic.get("blocking", []) or []
    if blocking:
        _add_subtitle(doc, "Items BLOQUANTS")
        for item in blocking:
            _add_box(
                doc, f"BLOQUANT #{item.get('num', '')}",
                [item.get("text", "")],
                bg_color=ASCEND_PALETTE["box_bg_warm"],
                title_color=ASCEND_PALETTE["accent_primary"],
            )

    warnings = critic.get("warnings", []) or []
    if warnings:
        _add_subtitle(doc, "Avertissements")
        for item in warnings:
            _add_box(
                doc, f"AVERTISSEMENT #{item.get('num', '')}",
                [item.get("text", "")],
                bg_color=ASCEND_PALETTE["box_bg_neutral"],
                title_color=ASCEND_PALETTE["accent_secondary"],
            )

    observations = critic.get("observations", []) or []
    if observations:
        _add_subtitle(doc, "Observations")
        for item in observations:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r_num = p.add_run(f"{item.get('num', '')}. ")
            _set_run(r_num, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
            r = p.add_run(item.get("text", ""))
            _set_run(r, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["title_primary"])

    _page_break(doc)


def _build_red_team(doc, ctx: BenchContext, chapter_num: int) -> None:
    rt = ctx.bench.get("red_team") or {}
    if not rt:
        return

    _add_chapter_title(doc, f"{chapter_num:02d}", "Rapport red-team")
    _add_spacer(doc)

    patterns = rt.get("patterns", []) or []
    if patterns:
        _add_subtitle(doc, "Patterns de fragilité communs")
        for i, pat in enumerate(patterns, start=1):
            p = doc.add_paragraph()
            r_num = p.add_run(f"{i}. ")
            _set_run(r_num, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
            r = p.add_run(pat)
            _set_run(r, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    verdicts = rt.get("verdicts", {}) or {}
    if verdicts:
        _add_subtitle(doc, "Verdicts par acteur")
        tbl = doc.add_table(rows=1, cols=2)
        _add_table_header_row(tbl.rows[0], ["Acteur", "Verdict"])
        actor_name_by_id = {a["id"]: a.get("name", a["id"]) for a in ctx.bench.get("actors", [])}
        for aid, verdict in verdicts.items():
            row = tbl.add_row()
            pal = verdict_palette(verdict)
            _add_table_data_row(row, [actor_name_by_id.get(aid, aid), verdict])
            v_cell = row.cells[1]
            _set_cell_shading(v_cell, pal["bg"])
            v_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(pal["fg"])
            v_cell.paragraphs[0].runs[0].font.bold = True
        _add_spacer(doc)

    msgs = rt.get("messages_to_comex", []) or []
    if msgs:
        _add_subtitle(doc, "Messages à nuancer devant le COMEX")
        for msg in msgs:
            _add_box(
                doc, "MESSAGE", [msg],
                bg_color=ASCEND_PALETTE["box_bg_neutral"],
                title_color=ASCEND_PALETTE["bleu_petrole"],
            )

    _page_break(doc)


def _build_transverse(doc, ctx: BenchContext, chapter_num: int) -> None:
    trans = ctx.bench.get("transverse") or {}
    if not trans:
        return

    _add_chapter_title(doc, f"{chapter_num:02d}", "Analyse transverse")
    _add_spacer(doc)

    patterns = trans.get("patterns_communs", []) or []
    if patterns:
        _add_subtitle(doc, "Patterns communs — ce que font ces acteurs pour s'émanciper")
        for i, pat in enumerate(patterns, start=1):
            p = doc.add_paragraph()
            r_num = p.add_run(f"{i}. ")
            _set_run(r_num, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
            r = p.add_run(pat)
            _set_run(r, font=FONT_BODY, size_pt=10,
                     color_hex=ASCEND_PALETTE["title_primary"])
        _add_spacer(doc)

    two_paths = trans.get("two_paths_table") or {}
    if two_paths and two_paths.get("dimensions"):
        _add_subtitle(doc, "Deux voies observées")
        dimensions = two_paths.get("dimensions", [])
        remediation = two_paths.get("remediation", [])
        by_design = two_paths.get("by_design", [])
        hybrid = two_paths.get("hybrid", [])
        ncols = 2 + (1 if hybrid else 0) + (1 if by_design else 0)
        headers = ["Dimension", "Remédiation"]
        if by_design:
            headers.append("By design")
        if hybrid:
            headers.append("Hybride")

        tbl = doc.add_table(rows=1, cols=ncols)
        _add_table_header_row(tbl.rows[0], headers)
        for i, dim in enumerate(dimensions):
            row = tbl.add_row()
            vals = [dim, remediation[i] if i < len(remediation) else ""]
            if by_design:
                vals.append(by_design[i] if i < len(by_design) else "")
            if hybrid:
                vals.append(hybrid[i] if i < len(hybrid) else "")
            _add_table_data_row(row, vals)
        _add_spacer(doc)

    ai_impact = trans.get("ai_impact", "")
    if ai_impact:
        _add_subtitle(doc, "Impact de l'IA sur la résilience numérique")
        _add_styled_para(doc, ai_impact, font=FONT_BODY, size=10,
                         color=ASCEND_PALETTE["title_primary"])

    _page_break(doc)


def _build_recommendations(doc, ctx: BenchContext, chapter_num: int) -> None:
    recs = ctx.bench.get("recommendations", []) or []
    if not recs:
        return

    _add_chapter_title(doc, f"{chapter_num:02d}", f"Enseignements pour {ctx.client()}")
    _add_spacer(doc)

    for rec in recs:
        # Bloc recommandation
        lines = [rec.get("body", "")]
        _add_box(
            doc, f"{rec.get('id', '')} — {rec.get('title', '')}",
            lines,
            bg_color=ASCEND_PALETTE["box_bg_neutral"],
            title_color=ASCEND_PALETTE["accent_primary"],
        )
        if rec.get("caveats"):
            _add_box(
                doc, "ATTENTION / CAVEAT",
                [rec["caveats"]],
                bg_color=ASCEND_PALETTE["box_bg_warm"],
                title_color=ASCEND_PALETTE["accent_primary"],
            )

    _page_break(doc)


def _build_exclusions_dealbreaker(doc, ctx: BenchContext, chapter_num: int) -> bool:
    """Section acteurs écartés par dealbreaker (v1.3.1).

    Conditionnelle sur exec_summary.exclus_dealbreaker non vide.
    Retourne True si la section a été écrite (pour incrémenter chapter_num),
    False sinon.
    """
    excluded_ids = (ctx.bench.get("exec_summary") or {}).get("exclus_dealbreaker") or []
    if not excluded_ids:
        return False

    actor_by_id = {a["id"]: a for a in ctx.bench.get("actors", [])}
    excluded_actors = [actor_by_id[aid] for aid in excluded_ids if aid in actor_by_id]
    if not excluded_actors:
        return False

    _add_chapter_title(doc, f"{chapter_num:02d}", "Exclusions — dealbreakers")
    _add_spacer(doc)
    _add_styled_para(
        doc,
        "Acteurs ecartes lors du check Phase 0 avant scoring ponderation. "
        "Chaque violation est lieee a une regle du scope.yaml.dealbreakers, "
        "avec preuve factuelle sourcee.",
        font=FONT_BODY, size=10, color=ASCEND_PALETTE["title_primary"],
    )
    _add_spacer(doc)

    for actor in excluded_actors:
        violations = actor.get("dealbreaker_violations", []) or []
        lines = []
        if violations:
            for v in violations:
                rule = v.get("rule_id", "?")
                evidence = v.get("evidence", "(pas de preuve documentee)")
                lines.append(f"{rule}: {evidence}")
        else:
            lines.append("Verdict violated sans detail de violations (dealbreaker_violations vide)")
        _add_box(
            doc, f"{actor.get('name', actor['id'])} — EXCLU",
            lines,
            bg_color=ASCEND_PALETTE["box_bg_warm"],
            title_color=ASCEND_PALETTE["accent_primary"],
        )

    _page_break(doc)
    return True


def _build_sources_annex(doc, ctx: BenchContext, chapter_num: int) -> None:
    _add_chapter_title(doc, f"{chapter_num:02d}", "Annexe — sources")
    _add_spacer(doc)

    # Grouper par actor_ref (ordre de rang)
    rank_by_actor_id = {a["id"]: a.get("rank", 999) for a in ctx.actors_by_rank}
    actor_name_by_id = {a["id"]: a.get("name", a["id"]) for a in ctx.bench.get("actors", [])}

    def src_key(s: dict):
        ar = s.get("actor_ref", "global")
        return (rank_by_actor_id.get(ar, 999), ar, s.get("id", ""))

    sorted_sources = sorted(ctx.bench.get("sources", []), key=src_key)

    current_actor = None
    for s in sorted_sources:
        ar = s.get("actor_ref", "global")
        if ar != current_actor:
            current_actor = ar
            label = actor_name_by_id.get(ar, ar)
            _add_subtitle(doc, label)

        p = doc.add_paragraph()
        r_id = p.add_run(f"[{s.get('id', '')}] ")
        _set_run(r_id, font=FONT_BODY, size_pt=9,
                 color_hex=ASCEND_PALETTE["accent_primary"], bold=True)
        r_title = p.add_run(f"{s.get('title', '')} — ")
        _set_run(r_title, font=FONT_BODY, size_pt=9,
                 color_hex=ASCEND_PALETTE["title_primary"], bold=True)
        r_meta = p.add_run(
            f"{s.get('author', '')} ({s.get('date', '')}) — "
            f"{s.get('type', '')} — fiabilité {s.get('reliability', '-')}/5"
        )
        _set_run(r_meta, font=FONT_BODY, size_pt=9,
                 color_hex=ASCEND_PALETTE["brun_clair"])
        if s.get("url"):
            p_url = doc.add_paragraph()
            p_url.paragraph_format.left_indent = Cm(0.5)
            r_url = p_url.add_run(s["url"])
            _set_run(r_url, font=FONT_BODY, size_pt=8,
                     color_hex=ASCEND_PALETTE["bleu_petrole"], italic=True)


# -----------------------------------------------------------------------------
# Point d'entrée
# -----------------------------------------------------------------------------


def write_docx(ctx: BenchContext, out_path: Path) -> None:
    doc = Document()
    _setup_sections(doc, ctx)

    _build_cover(doc, ctx)
    _build_toc(doc, ctx)
    _build_how_to_read(doc, ctx)
    _build_exec_summary(doc, ctx)
    _build_ranking(doc, ctx)

    next_chapter = 4
    for actor in ctx.actors_by_rank:
        _build_actor_sheet(doc, ctx, actor, next_chapter)
        next_chapter += 1

    _build_critic(doc, ctx, next_chapter); next_chapter += 1
    _build_red_team(doc, ctx, next_chapter); next_chapter += 1
    _build_transverse(doc, ctx, next_chapter); next_chapter += 1
    _build_recommendations(doc, ctx, next_chapter); next_chapter += 1
    # v1.3.1 : section Exclusions dealbreaker si non vide, incrémente chapter_num
    if _build_exclusions_dealbreaker(doc, ctx, next_chapter):
        next_chapter += 1
    _build_sources_annex(doc, ctx, next_chapter)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
